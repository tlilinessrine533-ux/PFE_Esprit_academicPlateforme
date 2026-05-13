from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT_PATH = r"c:\Users\Asus\Desktop\NV PFE\Rapport_PFE_Plateforme_Academique_ESPRIT.docx"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m_name, m_val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m_name}"))
        if node is None:
            node = OxmlElement(f"w:{m_name}")
            tcMar.append(node)
        node.set(qn("w:w"), str(m_val))
        node.set(qn("w:type"), "dxa")


def add_toc(paragraph):
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = "Cliquez-droit puis Mettre a jour le champ pour generer le sommaire automatique."
    run.append(text)
    fld_simple.append(run)
    paragraph._p.append(fld_simple)


def add_heading_paragraph(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows, style="Table Grid"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(hdr_cells[i])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])

    doc.add_paragraph("")
    return table


def set_global_style(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for hname, size in [("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)]:
        hstyle = doc.styles[hname]
        hstyle.font.name = "Calibri"
        hstyle.font.size = Pt(size)
        hstyle.font.bold = True
        hstyle.font.color.rgb = RGBColor(31, 78, 121)


def build_document():
    doc = Document()
    set_global_style(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("REPUBLIQUE TUNISIENNE")
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MINISTERE DE L'ENSEIGNEMENT SUPERIEUR ET DE LA RECHERCHE SCIENTIFIQUE")
    r.bold = True
    r.font.size = Pt(12)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RAPPORT DE PROJET DE FIN D'ETUDES")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Conception et developpement d'une plateforme intelligente de suivi des activites academiques")
    r.bold = True
    r.font.size = Pt(15)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Plateforme academique ESPRIT")
    r.font.size = Pt(13)

    doc.add_paragraph("")
    add_table(
        doc,
        ["Rubrique", "Contenu"],
        [
            ["Etudiant(e)", "[A completer]"],
            ["Encadrant academique", "[A completer]"],
            ["Organisme d'accueil", "ESPRIT"],
            ["Annee universitaire", "2025-2026"],
            ["Date de version", "5 mai 2026"],
        ],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Document de reference - version pre-soutenance").italic = True

    doc.add_page_break()

    # TOC
    add_heading_paragraph(doc, "Sommaire", level=1)
    toc_p = doc.add_paragraph()
    add_toc(toc_p)
    doc.add_page_break()

    # General introduction
    add_heading_paragraph(doc, "Introduction generale", level=1)
    add_body(
        doc,
        "La transformation digitale des etablissements d'enseignement superieur impose la mise en place de systemes "
        "capables de centraliser les donnees academiques, fluidifier les processus de validation et produire des indicateurs "
        "fiables pour la prise de decision. Dans cette perspective, ce projet propose une plateforme web intelligente qui "
        "oriente les flux metier autour d'un workflow tracable, d'un pilotage par KPI et d'un reporting multi-niveaux."
    )
    add_body(
        doc,
        "Ce rapport presente le cadrage du projet, la specification des besoins avec la mise en oeuvre Scrum et l'approche "
        "decisionnelle retenue pour l'analyse des performances academiques. Les donnees de reference exploitees dans ce document "
        "sont issues d'une extraction du 5 mai 2026 sur la base pfe_academic_platform."
    )

    # Chapter 1
    add_heading_paragraph(doc, "Chapitre 1 : Cadre general du projet", level=1)
    add_heading_paragraph(doc, "Introduction", level=2)
    add_body(
        doc,
        "Ce chapitre situe le projet dans son environnement institutionnel, formalise la problematique et justifie la "
        "demarche de realisation."
    )

    add_heading_paragraph(doc, "I. Presentation de l'organisme d'accueil", level=2)
    add_body(
        doc,
        "L'organisme d'accueil est un etablissement d'enseignement superieur structure autour de plusieurs departements "
        "academiques, avec une gouvernance basee sur des roles distincts : enseignant, chef de departement, administration "
        "et super administration."
    )
    add_body(
        doc,
        "Ses missions principales incluent la qualite pedagogique, la production scientifique, la conformite administrative "
        "et l'evaluation des performances individuelles et collectives. Ces missions necessitent un systeme d'information "
        "capable d'assurer la coherence des donnees et la transparence des decisions."
    )

    add_heading_paragraph(doc, "II. Presentation du projet", level=2)

    add_heading_paragraph(doc, "1. Analyse de l'existant", level=3)
    add_body(
        doc,
        "L'analyse preliminaire a mis en evidence un fonctionnement fragmente : multiplicite de fichiers, consolidations "
        "manuelles, heterogeneite des formats et manque de synchronisation des etats de validation."
    )

    add_heading_paragraph(doc, "2. Critique de l'existant", level=3)
    add_bullets(doc, [
        "Absence d'une source de verite unique pour les activites academiques.",
        "Tracabilite insuffisante des decisions de validation et des commentaires metier.",
        "Charge administrative elevee lors des consolidations periodiques.",
        "Faible capacite de pilotage comparatif entre enseignants et departements.",
        "Reporting lent et faiblement standardise."
    ])

    add_heading_paragraph(doc, "3. Problematique", level=3)
    add_body(
        doc,
        "Comment concevoir une plateforme securisee, evolutive et orientee decisionnel permettant de centraliser les "
        "activites academiques, d'automatiser les workflows de validation et de produire des indicateurs de performance "
        "fiables pour le pilotage institutionnel ?"
    )

    add_heading_paragraph(doc, "4. Solution proposee", level=3)
    add_body(
        doc,
        "La solution realisee est une architecture web trois couches : frontend Angular 21, backend Spring Boot (Java 21) "
        "et base MySQL, avec API REST securisee (JWT + mecanismes renforces) et modules metier couvrant le cycle "
        "declaration-validation-reporting."
    )
    add_table(
        doc,
        ["Axe", "Implementation"],
        [
            ["Frontend", "Angular 21 + TypeScript + routing protege par roles"],
            ["Backend", "Spring Boot 3.4.4, services metier, DTO, controleurs REST"],
            ["Securite", "JWT, 2FA, passkeys, reset mot de passe, restrictions RBAC"],
            ["Donnees", "MySQL avec modele d'activites et historique de validation"],
            ["Decisionnel", "Dashboards personnel/departement/global + benchmarks + trends"],
            ["Reporting", "PDF/Excel individuel, departemental et institutionnel"],
        ],
    )

    add_heading_paragraph(doc, "III. Methodologie de gestion de projet", level=2)
    add_heading_paragraph(doc, "1. Comparaison entre les methodologies", level=3)
    add_table(
        doc,
        ["Critere", "Cycle en V", "Cascade", "Agile/Scrum"],
        [
            ["Gestion du changement", "Faible", "Tres faible", "Forte"],
            ["Livraisons incrementales", "Limitees", "Rares", "Regulieres"],
            ["Feedback utilisateur", "Tardif", "Tres tardif", "Continu"],
            ["Adaptation a un produit evolutif", "Moyenne", "Faible", "Tres adaptee"],
            ["Risque de decalage besoin/produit", "Eleve", "Eleve", "Reduit"],
        ],
    )

    add_heading_paragraph(doc, "2. Choix de la methodologie", level=3)
    add_body(
        doc,
        "La demarche Agile a ete retenue pour sa capacite d'adaptation, son orientation valeur et sa compatibilite avec "
        "un produit a forte composante metier evolutive."
    )

    add_heading_paragraph(doc, "3. Choix de la methode", level=3)
    add_body(
        doc,
        "La methode Scrum a ete adoptee pour piloter le projet par increments : backlog priorise, sprints focalises, "
        "demonstrations intermediaires et ajustements rapides."
    )

    add_heading_paragraph(doc, "Conclusion", level=2)
    add_body(
        doc,
        "Le cadrage confirme la pertinence d'une plateforme academique unifiee et d'une gouvernance de projet agile."
    )

    # Chapter 2
    add_heading_paragraph(doc, "Chapitre 2 : Specification des besoins et mise en place avec SCRUM", level=1)
    add_heading_paragraph(doc, "Introduction", level=2)
    add_body(
        doc,
        "Ce chapitre formalise les besoins metier, les exigences non fonctionnelles, les KPI cibles et l'organisation Scrum."
    )

    add_heading_paragraph(doc, "I. Analyse des besoins", level=2)

    add_heading_paragraph(doc, "1. Identification des acteurs", level=3)
    add_table(
        doc,
        ["Acteur", "Responsabilites principales"],
        [
            ["Enseignant", "Declarer ses activites, soumettre, suivre les statuts, generer ses rapports"],
            ["Chef de departement", "Traiter les demandes du departement, valider/rejeter/corriger"],
            ["Administration", "Validation finale, evaluations, parametres de bonus/points, pilotage global"],
            ["Super administrateur", "Gouvernance des acces, supervision securite et parametres critiques"],
        ],
    )

    add_heading_paragraph(doc, "2. Les besoins fonctionnels", level=3)
    add_bullets(doc, [
        "Authentification securisee et gestion des profils.",
        "Gestion complete des utilisateurs et des roles.",
        "Gestion des activites : enseignement, encadrement, recherche, evenement, surveillance, responsabilite, disponibilite.",
        "Workflow complet : soumission, validation departementale, validation finale, rejet et correction.",
        "Tableaux de bord dynamiques (personnel, departemental, global).",
        "Generation de rapports PDF et Excel avec historique de telechargement.",
        "Notifications internes et emails sur les etapes critiques.",
        "Evaluation administrative des enseignants avec calcul de bonus et decision finale."
    ])

    add_heading_paragraph(doc, "3. Les besoins non-fonctionnels", level=3)
    add_table(
        doc,
        ["Categorie", "Exigence"],
        [
            ["Securite", "JWT, controle d'acces par role, 2FA/passkey, tracabilite des operations"],
            ["Performance", "Reponses rapides sur API metier et dashboards"],
            ["Fiabilite", "Cohesion transactionnelle et gestion robuste des erreurs"],
            ["Maintenabilite", "Architecture modulaire et separation claire des responsabilites"],
            ["Evolutivite", "Ajout simple de types d'activites, KPI et rapports"],
            ["Auditabilite", "Historique detaille de validation et journalisation des decisions"],
        ],
    )

    add_heading_paragraph(doc, "4. Identifications des KPIs", level=3)
    add_table(
        doc,
        ["KPI", "Definition", "Source"],
        [
            ["Taux de validation", "Activites validees / activites declarees", "activities.status"],
            ["Heures realisees", "Somme des heures d'enseignement effectives", "teaching_activities.completed_hours"],
            ["Production scientifique", "Nombre de publications/recherches et indexation", "research_activities"],
            ["Encadrement actif", "Nombre d'encadrements par statut", "supervision_activities"],
            ["Activites en attente", "Dossiers soumis non encore traites", "workflow"],
            ["Rapports generes", "Volume des exports par type et format", "reports"],
            ["Bonus estime", "Points cumules - penalites absences", "administration settings"],
        ],
    )

    add_heading_paragraph(doc, "II. Mise en place du projet avec SCRUM", level=2)

    add_heading_paragraph(doc, "1. Distribution des roles", level=3)
    add_table(
        doc,
        ["Role Scrum", "Affectation", "Mission"],
        [
            ["Product Owner", "Encadrement metier", "Prioriser la valeur et valider les increments"],
            ["Scrum Master", "Equipe de realisation", "Assurer la cadence et lever les blocages"],
            ["Development Team", "Equipe technique", "Concevoir, developper, tester et documenter"],
        ],
    )

    add_heading_paragraph(doc, "2. Backlog Product", level=3)
    add_table(
        doc,
        ["ID", "User Story", "Priorite", "Estimation"],
        [
            ["US01", "Authentification securisee", "Haute", "8 SP"],
            ["US02", "Gestion des utilisateurs et roles", "Haute", "8 SP"],
            ["US03", "Declaration des enseignements", "Haute", "8 SP"],
            ["US04", "Declaration des encadrements", "Haute", "5 SP"],
            ["US05", "Declaration des recherches", "Haute", "5 SP"],
            ["US06", "Workflow de validation multi-niveaux", "Haute", "8 SP"],
            ["US07", "Dashboards multi-perimetres", "Haute", "8 SP"],
            ["US08", "Reporting PDF/Excel", "Haute", "8 SP"],
            ["US09", "Notifications et relances", "Moyenne", "3 SP"],
            ["US10", "Evaluation administrative et bonus", "Moyenne", "5 SP"],
        ],
    )

    add_heading_paragraph(doc, "3. Sprint Backlog", level=3)
    add_table(
        doc,
        ["Sprint", "Objectif", "Livrables"],
        [
            ["Sprint 0", "Cadrage et socle technique", "Architecture, schema BD, environnement"],
            ["Sprint 1", "Securite et comptes", "Auth JWT, profils, gestion users"],
            ["Sprint 2", "Enseignement + workflow", "Soumission/validation/historique"],
            ["Sprint 3", "Extensions activites", "Encadrement, recherche, evenement, surveillance"],
            ["Sprint 4", "Dashboards", "KPI personnels, departementaux, globaux"],
            ["Sprint 5", "Reporting et stabilisation", "PDF/Excel, corrections, documentation"],
        ],
    )

    add_heading_paragraph(doc, "4. Diagramme de GANT", level=3)
    add_table(
        doc,
        ["Periode", "Activites"],
        [
            ["S1-S2", "Cadrage, backlog final, modelisation"],
            ["S3-S4", "Authentification et gestion des acces"],
            ["S5-S7", "Modules activites coeur"],
            ["S8-S9", "Workflow et historique"],
            ["S10-S12", "Dashboards et KPIs"],
            ["S13-S14", "Reporting PDF/Excel"],
            ["S15", "Tests et correction"],
            ["S16", "Documentation finale et soutenance"],
        ],
    )

    add_heading_paragraph(doc, "Conclusion", level=2)
    add_body(
        doc,
        "La mise en oeuvre Scrum a permis une progression incrementale maitrisee, avec des livrables metier evaluables "
        "a chaque sprint."
    )

    # Chapter 3
    add_heading_paragraph(doc, "Chapitre 3 : Analyse des donnees et conception multidimensionnelle", level=1)
    add_heading_paragraph(doc, "Introduction", level=2)
    add_body(
        doc,
        "Ce chapitre presente le volet decisionnel du projet. Il s'appuie sur une extraction de la base en date du 5 mai 2026 "
        "et propose une architecture d'analyse orientee performance academique."
    )

    add_heading_paragraph(doc, "1. Sprint Backlog decisionnel", level=2)
    add_table(
        doc,
        ["Sprint decisionnel", "Objectif"],
        [
            ["SD1", "Audit de qualite des donnees et stabilisation des indicateurs"],
            ["SD2", "Analyse exploratoire univariee et bivariee"],
            ["SD3", "Conception multidimensionnelle et schema final"],
        ],
    )

    add_heading_paragraph(doc, "2. Analyse exploratoire des donnees", level=2)
    add_heading_paragraph(doc, "2.1 Presentation des donnees", level=3)
    add_table(
        doc,
        ["Entite", "Volume observe"],
        [
            ["Utilisateurs", "76"],
            ["Departements", "7"],
            ["Activites totales", "4989"],
            ["Enseignements", "1389"],
            ["Encadrements", "881"],
            ["Recherches", "660"],
            ["Evenements", "659"],
            ["Surveillances", "877"],
            ["Responsabilites", "513"],
            ["Demandes disponibilite", "10"],
            ["Historique validations", "2142"],
            ["Rapports", "19"],
            ["Notifications", "62"],
        ],
    )

    add_body(
        doc,
        "Repartition des roles : 73 enseignants, 1 chef de departement, 1 administration, 1 super administrateur. "
        "La population enseignante est composee de 53 permanents et 20 vacataires."
    )

    add_heading_paragraph(doc, "2.2 Analyse univariee", level=3)
    add_table(
        doc,
        ["Type d'activite (2025-2026)", "Nombre", "Pourcentage"],
        [
            ["Enseignement", "322", "26,37 %"],
            ["Encadrement", "224", "18,35 %"],
            ["Surveillance", "220", "18,02 %"],
            ["Recherche", "149", "12,20 %"],
            ["Responsabilite", "148", "12,12 %"],
            ["Evenement", "148", "12,12 %"],
            ["Disponibilite", "10", "0,82 %"],
        ],
    )

    add_table(
        doc,
        ["Statut (2025-2026)", "Nombre", "Pourcentage"],
        [
            ["VALIDEE_FINALE", "579", "47,42 %"],
            ["VALIDEE_DEPARTEMENT", "269", "22,03 %"],
            ["SOUMISE", "207", "16,95 %"],
            ["A_CORRIGER", "72", "5,90 %"],
            ["REJETEE", "66", "5,41 %"],
            ["BROUILLON", "28", "2,29 %"],
        ],
    )

    add_heading_paragraph(doc, "2.3 Analyse bivariée", level=3)
    add_table(
        doc,
        ["Couple d'analyse", "Observation"],
        [
            ["Type enseignant vs heures d'enseignement", "Permanents : 7979,50 h cumulees; Vacataires : 0 h (jeu de demo sans points enseignement)."],
            ["Type activite vs taux de validation", "Responsabilite 89,86 % ; Enseignement 69,57 % ; Encadrement 61,16 %."],
            ["Partenariat vs points moyens enseignement", "Sans partenariat: 22,88 pts ; Academique: 43,23 pts ; Professionnelle: 44,23 pts."],
            ["Departement vs taux de validation", "Variation observee de 64,20 % a 72,39 % selon le departement."],
            ["Annee academique vs taux global", "2022-2023: 100 % ; 2023-2024: 100 % ; 2024-2025: 96,20 % ; 2025-2026: 69,45 %."],
        ],
    )

    add_heading_paragraph(doc, "3. Conception multidimensionnelle", level=2)

    add_heading_paragraph(doc, "3.1 Choix de l'approche", level=3)
    add_heading_paragraph(doc, "3.1.1 Approche Kimball", level=3)
    add_body(doc, "Approche orientee datamarts metier, adaptee aux besoins analytiques iteratifs et a la livraison agile.")
    add_heading_paragraph(doc, "3.1.2 Approche Inmon", level=3)
    add_body(doc, "Approche top-down centree entrepot d'entreprise, robuste mais plus lourde pour un cycle PFE.")
    add_heading_paragraph(doc, "3.1.3 Comparaison et choix", level=3)
    add_body(doc, "Le choix retenu est l'approche Kimball pour privilegier la rapidite de mise en oeuvre, la lisibilite des KPI et la flexibilite d'evolution.")

    add_heading_paragraph(doc, "3.2 Choix du schema conceptuel", level=3)
    add_heading_paragraph(doc, "3.2.1 Schema en etoile", level=3)
    add_body(doc, "Simple et performant, mais limite pour couvrir plusieurs processus complexes inter-relies.")
    add_heading_paragraph(doc, "3.2.2 Schema en flocon de neige", level=3)
    add_body(doc, "Plus normalise, meilleur controle de redondance, mais requetes analytiques plus complexes.")
    add_heading_paragraph(doc, "3.2.3 Schema en constellation", level=3)
    add_body(doc, "Plusieurs faits partagent des dimensions conformes ; tres adapte a un SI academique multi-processus.")
    add_heading_paragraph(doc, "3.2.4 Comparaison et choix", level=3)
    add_body(doc, "Le schema retenu est la constellation, afin d'integrer workflow, activites, reporting et evaluations administratives.")

    add_heading_paragraph(doc, "3.3 Presentation des tables de dimensions", level=3)
    add_table(
        doc,
        ["Dimension", "Role analytique"],
        [
            ["DimDate", "Analyse par jour, mois, trimestre, annee academique, semestre"],
            ["DimEnseignant", "Profil enseignant, role, type, statut"],
            ["DimDepartement", "Pilotage par structure academique"],
            ["DimTypeActivite", "Classification metier des activites"],
            ["DimStatut", "Cycle de vie des dossiers"],
            ["DimValidation", "Niveau et decision de validation"],
            ["DimRapport", "Type, format et perimetre des rapports"],
            ["DimContexteEnseignement", "Semestre, mode, langue, partenariat"],
        ],
    )

    add_heading_paragraph(doc, "3.4 Presentation des tables de faits", level=3)
    add_table(
        doc,
        ["Table de faits", "Grain", "Mesures"],
        [
            ["FactActivite", "1 ligne = 1 activite", "Heures, points, volume, statut"],
            ["FactValidationWorkflow", "1 ligne = 1 decision", "Delai de traitement, nature de decision"],
            ["FactReporting", "1 ligne = 1 rapport", "Nb rapports, type, format, periode"],
            ["FactEvaluationAdministrative", "1 ligne = 1 evaluation enseignant/periode", "Points, bonus, absences, decision finale"],
        ],
    )

    add_heading_paragraph(doc, "3.5 Schema final", level=3)
    add_body(
        doc,
        "Le schema final propose une constellation decisionnelle avec dimensions conformes partagees entre les faits. "
        "Cette structure garantit la coherence des calculs KPI, la comparaison inter-departements et l'analyse temporelle longitudinale."
    )

    add_heading_paragraph(doc, "Conclusion", level=2)
    add_body(
        doc,
        "L'analyse exploratoire confirme la maturite des donnees et la pertinence d'un modele multidimensionnel en constellation "
        "pour soutenir le pilotage academique."
    )

    # General conclusion
    add_heading_paragraph(doc, "Conclusion generale", level=1)
    add_body(
        doc,
        "Le projet realise une transformation operationnelle et decisionnelle du suivi academique. Il unifie la collecte des "
        "activites, fiabilise les validations, automatise les reportings et offre un cadre d'evaluation objectivable des performances."
    )
    add_body(
        doc,
        "Sur le plan technique, l'architecture Angular/Spring Boot/MySQL offre un socle robuste et evolutif. Sur le plan metier, "
        "la plateforme renforce la transparence, la responsabilisation et la capacite de pilotage a tous les niveaux de gouvernance."
    )

    # Recommendations
    add_heading_paragraph(doc, "Perspectives et recommandations", level=1)
    add_numbered(doc, [
        "Generaliser le workflow complet a tous les types d'activites avec regles metier homogenes.",
        "Industrialiser un entrepot decisionnel dedie (datamarts par domaine).",
        "Mettre en place des tests automatises de non-regression KPI.",
        "Ajouter des tableaux de bord predictifs (charge, performance, risque de retard).",
        "Renforcer la gouvernance des donnees (qualite, lineage, data catalog).",
    ])

    # Bibliography
    add_heading_paragraph(doc, "Bibliographie", level=1)
    add_numbered(doc, [
        "Documentation technique interne du projet : DOCUMENTATION_TECHNIQUE_PFE.md.",
        "Manuel utilisateur interne : MANUEL_UTILISATEUR_PFE.md.",
        "Schema relationnel et scripts SQL : schema-pfe-mysql.sql, donnees-test-pfe.sql.",
        "Guides d'implementation backend/frontend : GUIDE_BACKEND_ETAPE_*.md, GUIDE_FRONTEND_ANGULAR_ETAPE_1.md.",
        "Spring Boot Reference Documentation.",
        "Angular Documentation.",
        "OWASP Application Security Verification Standard (ASVS).",
        "Scrum Guide 2020, Ken Schwaber & Jeff Sutherland.",
        "Kimball R., Ross M. - The Data Warehouse Toolkit.",
        "Inmon W. H. - Building the Data Warehouse.",
    ])

    # Appendices
    add_heading_paragraph(doc, "Annexes", level=1)

    add_heading_paragraph(doc, "Annexe A - Synthese technique du socle", level=2)
    add_table(
        doc,
        ["Element", "Valeur"],
        [
            ["Backend", "Spring Boot 3.4.4 - Java 21"],
            ["Frontend", "Angular 21.2.x"],
            ["Base de donnees", "MySQL"],
            ["Securite", "JWT + 2FA + passkey + controle RBAC"],
            ["Nombre de modules backend", "15 packages metier principaux"],
            ["Nombre d'endpoints exposes", "107 endpoints"],
            ["Proxy frontend", "ng serve --proxy-config proxy.conf.json"],
        ],
    )

    add_heading_paragraph(doc, "Annexe B - Dictionnaire de donnees (extrait)", level=2)
    add_table(
        doc,
        ["Table", "Description"],
        [
            ["users", "Comptes utilisateurs, roles, rattachement departement"],
            ["activities", "Entite mere des activites academiques"],
            ["teaching_activities", "Details des activites d'enseignement"],
            ["supervision_activities", "Details des encadrements et jurys"],
            ["research_activities", "Details des activites de recherche"],
            ["event_activities", "Details des evenements scientifiques"],
            ["exam_surveillance_activities", "Details des surveillances d'examens"],
            ["responsibility_activities", "Responsabilites academiques"],
            ["availability_request_activities", "Demandes de conge/mission"],
            ["validation_history", "Trace complete des decisions de validation"],
            ["reports", "Historique des rapports generes"],
            ["notifications", "Notifications utilisateur"],
        ],
    )

    add_heading_paragraph(doc, "Annexe C - Catalogue des KPI (extrait)", level=2)
    add_table(
        doc,
        ["KPI", "Formule de calcul"],
        [
            ["Taux de validation", "(Nb validees departement + final) / Nb total activites"],
            ["Ecart horaire", "Heures realisees - Heures planifiees"],
            ["Points enseignement", "Somme des composantes pedagogiques validees"],
            ["Production recherche", "Nb publications + indexations + points recherche"],
            ["Dynamique workflow", "Nb soumis, a corriger, rejetes, valides"],
            ["Bonus estime", "Base + (points x coefficient) - penalites absences"],
        ],
    )

    add_heading_paragraph(doc, "Annexe D - KPI observes (extraction du 5 mai 2026)", level=2)
    add_table(
        doc,
        ["Indicateur", "Valeur"],
        [
            ["Activites totales (toutes annees)", "4989"],
            ["Activites 2025-2026", "1221"],
            ["Taux de validation global 2025-2026", "69,45 %"],
            ["Cours 2025-2026", "322"],
            ["Encadrements 2025-2026", "224"],
            ["Recherches 2025-2026", "149"],
            ["Rapports generes", "19"],
            ["Decisions workflow historisees", "2142"],
        ],
    )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_PATH)

